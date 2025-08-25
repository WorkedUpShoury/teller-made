# --- Imports ---
import os
import io
import spacy
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

# Try to import DB helper; provide a safe fallback if missing
try:
    from db import get_db_connection  # make sure backend/db.py defines this
except Exception:
    def get_db_connection():
        raise RuntimeError("get_db_connection() not available. Create backend/db.py with a get_db_connection implementation.")

# --- Configuration ---
app = Flask(__name__)
# Enable Cross-Origin Resource Sharing to allow your React app to communicate with this server
CORS(
    app,
    resources={r"/*": {"origins": [
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:3001",
        "http://127.0.0.1:3001",
        "http://192.168.1.104:3001",
    ]}},
    supports_credentials=True,
    expose_headers=["Content-Disposition"],
)

# Set up a folder for temporary file uploads
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# --- NLP and AI Setup ---
# Load spaCy model (will download if not present)
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    print("Downloading spaCy model 'en_core_web_sm'...")
    # Fix: import from public path instead of spacy.cli
    from spacy.cli.download import download  # <-- public import
    download('en_core_web_sm')
    nlp = spacy.load('en_core_web_sm')

# Configure Google Gemini API Key
# IMPORTANT: Set your Google API key in an environment variable named 'GEMINI_API_KEY'
try:
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
except Exception as e:
    print(f"Error configuring Google API: {e}")

# --- Helper Functions ---

def parse_document(file_path: str) -> str:
    """Parses a .pdf or .docx file and returns its text content."""
    text = ""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.pdf':
        with fitz.open(file_path) as doc:
            for page in doc:
                # Fix: be explicit about extractor to satisfy type checkers
                text += page.get_text("text")
    elif file_extension == '.docx':
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + '\n'
    else:
        raise ValueError("Unsupported file format. Please use PDF or DOCX.")
    return text

def extract_keywords_from_jd(job_desc: str):
    """Extracts key skills and nouns from a job description using spaCy."""
    doc = nlp(job_desc)
    keywords = set()
    # Using noun chunks for more meaningful phrases (e.g., "project management", "data analysis")
    for chunk in doc.noun_chunks:
        if chunk and chunk.text:
            keywords.add(chunk.text.lower().strip())
    # Using entities for specific technologies or organizations
    for ent in doc.ents:
        if ent.label_ in {"ORG", "PRODUCT", "TECH"}:
            if ent and ent.text:
                keywords.add(ent.text.lower().strip())
    return list(keywords)

def get_ai_optimization(resume_text: str, job_desc: str) -> str | None:
    """Uses Google Gemini to rewrite the resume based on a job description."""
    # Initialize the Gemini model
    # Use 'gemini-1.5-flash-latest' for speed and cost-effectiveness
    # Use 'gemini-1.5-pro-latest' for higher quality results
    try:
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
    except Exception as e:
        print(f"Error creating GenerativeModel: {e}")
        return None

    prompt = f"""
You are an expert career coach and resume writer. Your task is to optimize a resume for a specific job description.

Instructions:
1) Thoroughly analyze the original resume and the target job description.
2) Rewrite the resume to align it with the job description, focusing on the Experience section.
3) Naturally integrate relevant keywords and skills from the job description into the resume's experience bullet points.
4) Rephrase responsibilities as quantifiable achievements. Use strong action verbs and metrics (e.g., "Increased sales by 20%," "Managed a team of 5").
5) Do NOT invent new experiences. Enhance the existing content to better reflect the job requirements.
6) Return ONLY the complete, rewritten resume text. Maintain a clean, professional, and ATS-friendly format. Do not add any extra commentary.

Original Resume Text:
---
{resume_text}
---

Target Job Description:
---
{job_desc}
---

Optimized Resume Text (return only the full resume text below this line):
""".strip()

    try:
        # Fix: avoid Pylance complaint about genai.types.GenerationConfig
        response = model.generate_content(
            prompt,
            generation_config={"temperature": 0.6}
        )
        return (getattr(response, "text", None) or "").strip()
    except Exception as e:
        print(f"Error calling Google Gemini API: {e}")
        return None

def create_optimized_docx(optimized_text: str) -> io.BytesIO:
    """Creates a new .docx file from text and returns it as a bytes buffer."""
    doc = Document()
    doc.add_heading('Optimized Resume', level=1)
    for line in optimized_text.split('\n'):
        # This basic logic helps preserve some structure.
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- API Routes ---

@app.route('/api/optimize', methods=['POST'])
def optimize_resume_endpoint():
    """Main endpoint to handle resume optimization."""
    if 'resumeFile' not in request.files or 'jobDesc' not in request.form:
        return jsonify({"error": "Missing resume file or job description"}), 400

    resume_file = request.files['resumeFile']
    job_desc = request.form.get('jobDesc', '')

    if not resume_file or not job_desc.strip():
        return jsonify({"error": "Missing resume file or job description"}), 400

    # Fix: guard filename (can be None)
    original_filename = resume_file.filename or "resume"
    safe_name = secure_filename(str(original_filename))
    if not safe_name:
        safe_name = "resume"

    if not os.getenv("GEMINI_API_KEY"):
        return jsonify({"error": "Server is not configured for AI processing. Missing GEMINI_API_KEY."}), 503

    file_path = None
    try:
        # 1. Save and Parse Resume
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)
        resume_file.save(file_path)

        resume_text = parse_document(file_path)
        if not resume_text.strip():
            raise ValueError("Could not extract text from the resume file.")

        # 2. Get AI Optimization
        optimized_resume_text = get_ai_optimization(resume_text, job_desc)
        if not optimized_resume_text:
            return jsonify({"error": "Failed to generate optimized content via AI service."}), 500

        # 3. Create New DOCX in Memory
        doc_buffer = create_optimized_docx(optimized_resume_text)

        # 4. Send File for Download
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name='optimized_resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return jsonify({"error": "An internal server error occurred."}), 500
    finally:
        # 5. Clean Up Uploaded File
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass

@app.route('/api/register', methods=['POST'])
def register():
    """Register a new user. Requires db.get_db_connection() to be implemented."""
    data = request.get_json(silent=True) or {}
    username = data.get("username")
    email = data.get("email")
    password = data.get("password")  # TODO: hash before storing

    if not username or not email or not password:
        return jsonify({"error": "Missing fields"}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO users (username, email, password) VALUES (%s, %s, %s)",
            (username, email, password)
        )
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"message": "User registered successfully"})
    except Exception as e:
        print("Error in register:", e)
        return jsonify({"error": str(e)}), 500

# --- Main Execution ---
if __name__ == '__main__':
    # Running on port 5001 to avoid conflicts with React's default port 3000/3001
    app.run(debug=True, port=5001)
